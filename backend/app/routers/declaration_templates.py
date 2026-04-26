import io
import logging
import os
import uuid
from datetime import datetime, timezone
from typing import Optional

from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, File, HTTPException, Request, UploadFile, status
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, field_validator
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.auth.jwt import get_current_analyst, require_admin
from app.utils.rate_limit import limiter as _limiter
from app.config import Settings, get_settings
from app.database import get_db
from app.models.analyst import Analyst
from app.models.declaration_template import DeclarationTemplate
from app.utils.audit import log_audit
from app.utils.docx_utils import convert_docx_to_pdf_via_libreoffice
from app.utils.file_utils import sanitize_filename, content_disposition_filename

logger = logging.getLogger(__name__)

router = APIRouter(tags=["declaration-templates"])

VALID_PROVIDER_TYPES = {
    "correduria_seguros",
    "agencia_seguros",
    "colaborador_externo",
    "generador_leads",
}

VALID_ENTITY_TYPES = {"PF", "PJ"}

PROVIDER_TYPE_LABELS = {
    "correduria_seguros": "Correduría de Seguros",
    "agencia_seguros": "Agencia de Seguros",
    "colaborador_externo": "Colaborador Externo",
    "generador_leads": "Generador de Leads",
}

ENTITY_TYPE_LABELS = {
    "PJ": "Persona Jurídica",
    "PF": "Persona Física",
}

# Accepted MIME type for DOCX uploads
DOCX_MIME_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# Placeholder → partner_info field mapping per entity type.
# Format in DOCX: "[ CAMPO ]" (with spaces inside brackets).
PJ_PLACEHOLDER_MAP = {
    "[RAZÓN SOCIAL DE LA EMPRESA]": "razon_social",
    "[CIF DE LA EMPRESA]": "cif",
    "[DOMICILIO SOCIAL DE LA EMPRESA]": "domicilio_social",
    "[NOMBRE Y APELLIDOS DEL REPRESENTANTE LEGAL]": "nombre_representante",
    "[NIF DEL REPRESENTANTE LEGAL]": "nif_representante",
}

PF_PLACEHOLDER_MAP = {
    "[NOMBRE Y APELLIDOS]": "nombre_apellidos",
    "[NIF]": "nif",
    "[DOMICILIO]": "domicilio",
    # Signature line: some templates use the representative field, others repeat nombre_apellidos
    "[NOMBRE Y APELLIDOS DEL REPRESENTANTE LEGAL]": "nombre_apellidos",
}


# ---------------------------------------------------------------------------
# Pydantic schemas
# ---------------------------------------------------------------------------


class DeclarationTemplateInfo(BaseModel):
    provider_type: str
    entity_type: str
    provider_type_label: str
    entity_type_label: str
    original_filename: str
    uploaded_at: datetime
    uploaded_by_name: Optional[str] = None

    model_config = {"from_attributes": True}


class AllTemplatesResponse(BaseModel):
    templates: list[DeclarationTemplateInfo]


class GenerateRequest(BaseModel):
    partner_info: dict

    # S12: prevent any single field value from being excessively long
    @field_validator("partner_info")
    @classmethod
    def validate_partner_info_values(cls, v):
        for key, val in v.items():
            if isinstance(val, str) and len(val) > 1000:
                raise ValueError(f"Field '{key}' exceeds maximum length of 1000 characters")
        return v


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _get_template_dir(documents_base_path: str) -> str:
    template_dir = os.path.join(documents_base_path, "declaration_templates")
    os.makedirs(template_dir, exist_ok=True)
    return template_dir


# OOXML Word namespace — used to find <w:t> elements via lxml
_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _replace_in_paragraph_elem(p_elem, replacements: dict[str, str]) -> None:
    """
    Replace placeholder strings in a single paragraph XML element.

    Using lxml's findall on the raw XML element (para._p) guarantees we find
    ALL <w:t> nodes — including those nested inside <w:hyperlink>, <w:ins>
    (tracked changes), or other wrapper elements that para.runs would miss.

    Strategy:
      Pass 1 — try replacing within each individual <w:t> element. This preserves
               per-run formatting (bold, underline, etc.) because we never move text
               between runs. If a placeholder lives entirely inside one run, it is
               replaced in-place without touching any other run.
      Pass 2 — if any placeholders still remain (they span multiple runs), fall back
               to the "collapse all into first element" approach. This loses per-run
               formatting for the affected paragraph but is the only way to handle
               placeholders whose characters were split across run boundaries by Word.
    """
    w_t_elems = p_elem.findall(".//" + _W_NS + "t")
    if not w_t_elems:
        return

    # Pass 1: per-run replacement (preserves formatting)
    for elem in w_t_elems:
        if not elem.text:
            continue
        for placeholder, value in replacements.items():
            if placeholder in elem.text:
                elem.text = elem.text.replace(placeholder, value)

    # Pass 2: check for placeholders that span multiple runs
    full_text = "".join(elem.text or "" for elem in w_t_elems)
    remaining = {ph: v for ph, v in replacements.items() if ph in full_text}
    if not remaining:
        return

    # Fallback: collapse all text into the first element (loses per-run formatting,
    # but is the only way to handle placeholders split across run boundaries)
    new_text = full_text
    for placeholder, value in remaining.items():
        new_text = new_text.replace(placeholder, value)
    w_t_elems[0].text = new_text
    for elem in w_t_elems[1:]:
        elem.text = ""


def _replace_placeholders_in_docx(doc: DocxDocument, replacements: dict[str, str]) -> None:
    """Replace placeholder strings (e.g. '[ CAMPO ]') in a python-docx Document."""
    for para in doc.paragraphs:
        _replace_in_paragraph_elem(para._p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph_elem(para._p, replacements)


def _build_replacements(entity_type: str, partner_info: dict) -> dict[str, str]:
    """Build the full placeholder → value mapping for a given entity type."""
    today = datetime.now(timezone.utc)
    _MESES = [
        "", "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    ]
    date_replacements = {
        "[DÍA]": str(today.day),
        "[MES]": _MESES[today.month],
        "[AÑO]": str(today.year),
    }

    if entity_type == "PJ":
        field_map = PJ_PLACEHOLDER_MAP
    else:
        field_map = PF_PLACEHOLDER_MAP

    replacements: dict[str, str] = {}
    for placeholder, field_key in field_map.items():
        value = partner_info.get(field_key, "")
        replacements[placeholder] = str(value)

    replacements.update(date_replacements)
    return replacements


# ---------------------------------------------------------------------------
# GET /api/declaration-templates  (JWT required — list all for admin)
# ---------------------------------------------------------------------------


@router.get("/declaration-templates", response_model=AllTemplatesResponse)
async def list_all_templates(
    db: AsyncSession = Depends(get_db),
    current_analyst: Analyst = Depends(get_current_analyst),
):
    """
    Return all declaration templates (one per provider_type × entity_type combination).
    Used by the admin UI to show what has been uploaded.
    """
    result = await db.execute(
        select(DeclarationTemplate).options(
            selectinload(DeclarationTemplate.uploaded_by_analyst)
        )
    )
    templates = result.scalars().all()

    template_map = {
        (t.provider_type, t.entity_type): t for t in templates
    }

    items = []
    for pt in sorted(VALID_PROVIDER_TYPES):
        for et in ("PJ", "PF"):
            t = template_map.get((pt, et))
            if t:
                items.append(
                    DeclarationTemplateInfo(
                        provider_type=t.provider_type,
                        entity_type=t.entity_type,
                        provider_type_label=PROVIDER_TYPE_LABELS.get(t.provider_type, t.provider_type),
                        entity_type_label=ENTITY_TYPE_LABELS.get(t.entity_type, t.entity_type),
                        original_filename=t.original_filename,
                        uploaded_at=t.uploaded_at,
                        uploaded_by_name=(
                            t.uploaded_by_analyst.full_name
                            if t.uploaded_by_analyst
                            else None
                        ),
                    )
                )

    return AllTemplatesResponse(templates=items)


# ---------------------------------------------------------------------------
# GET /api/declaration-templates/{provider_type}/{entity_type}  (PUBLIC)
# ---------------------------------------------------------------------------


@router.get("/declaration-templates/{provider_type}/{entity_type}")
async def get_template_info(
    provider_type: str,
    entity_type: str,
    db: AsyncSession = Depends(get_db),
):
    """
    Public endpoint: returns metadata about a declaration template if it exists.
    Used by the partner invite page to decide whether to show a download button.
    Returns 404 if no template has been uploaded for this combination.
    """
    if provider_type not in VALID_PROVIDER_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Invalid provider_type. Must be one of: {', '.join(sorted(VALID_PROVIDER_TYPES))}",
        )
    if entity_type not in VALID_ENTITY_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="entity_type must be 'PF' or 'PJ'",
        )

    result = await db.execute(
        select(DeclarationTemplate).where(
            DeclarationTemplate.provider_type == provider_type,
            DeclarationTemplate.entity_type == entity_type,
        )
    )
    template = result.scalar_one_or_none()

    if template is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="No template uploaded")

    return {
        "provider_type": template.provider_type,
        "entity_type": template.entity_type,
        "original_filename": template.original_filename,
        "uploaded_at": template.uploaded_at.isoformat(),
    }


# ---------------------------------------------------------------------------
# GET /api/declaration-templates/{provider_type}/{entity_type}/download  (PUBLIC)
# ---------------------------------------------------------------------------


@router.get("/declaration-templates/{provider_type}/{entity_type}/download")
@_limiter.limit("60/hour")
async def download_template(
    request: Request,
    provider_type: str,
    entity_type: str,
    db: AsyncSession = Depends(get_db),
    settings: Settings = Depends(get_settings),
):
    """
    Public endpoint: download the raw DOCX template for a provider type + entity type.
    """
    if provider_type not in VALID_PROVIDER_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Invalid provider_type",
        )
    if entity_type not in VALID_ENTITY_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="entity_type must be 'PF' or 'PJ'",
        )

    result = await db.execute(
        select(DeclarationTemplate).where(
            DeclarationTemplate.provider_type == provider_type,
            DeclarationTemplate.entity_type == entity_type,
        )
    )
    template = result.scalar_one_or_none()

    if template is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="No template found")

    if not os.path.exists(template.file_path):
        logger.error(
            "Declaration template file missing on disk: %s (provider=%s entity=%s)",
            template.file_path,
            provider_type,
            entity_type,
        )
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Template file not found on disk",
        )

    return FileResponse(
        path=template.file_path,
        media_type=DOCX_MIME_TYPE,
        filename=template.original_filename,
        headers={
            "Content-Disposition": content_disposition_filename(template.original_filename),
        },
    )


# ---------------------------------------------------------------------------
# POST /api/declaration-templates/{provider_type}/{entity_type}/generate  (PUBLIC)
# ---------------------------------------------------------------------------


@router.post("/declaration-templates/{provider_type}/{entity_type}/generate")
@_limiter.limit("20/hour")
async def generate_declaration_pdf(
    request: Request,
    provider_type: str,
    entity_type: str,
    body: GenerateRequest,
    db: AsyncSession = Depends(get_db),
    settings: Settings = Depends(get_settings),
):
    """
    Public endpoint: fill placeholders in the DOCX template with partner data
    and return a personalised PDF.

    Pipeline:
      1. Load DOCX from disk
      2. Replace all [ PLACEHOLDER ] tokens with partner values (lxml approach)
      3. Convert patched DOCX → PDF via LibreOffice headless subprocess
      4. Stream the PDF back to the caller
    """
    if provider_type not in VALID_PROVIDER_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Invalid provider_type",
        )
    if entity_type not in VALID_ENTITY_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="entity_type must be 'PF' or 'PJ'",
        )

    result = await db.execute(
        select(DeclarationTemplate).where(
            DeclarationTemplate.provider_type == provider_type,
            DeclarationTemplate.entity_type == entity_type,
        )
    )
    template = result.scalar_one_or_none()

    if template is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No declaration template found for this provider type and entity type",
        )

    if not os.path.exists(template.file_path):
        logger.error(
            "Declaration template file missing on disk: %s", template.file_path
        )
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Template file not found on disk",
        )

    # ── 1. Load and patch the DOCX ────────────────────────────────────────────
    doc = DocxDocument(template.file_path)
    replacements = _build_replacements(entity_type, body.partner_info)
    _replace_placeholders_in_docx(doc, replacements)

    # ── 2. Save patched DOCX to an in-memory buffer ───────────────────────────
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()

    # ── 3. Convert patched DOCX → PDF via LibreOffice ────────────────────────
    pdf_bytes = await convert_docx_to_pdf_via_libreoffice(docx_bytes)

    # ── 4. Stream back ────────────────────────────────────────────────────────
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={
            "Content-Disposition": 'inline; filename="declaracion.pdf"',
        },
    )


# ---------------------------------------------------------------------------
# PUT /api/declaration-templates/{provider_type}/{entity_type}  (JWT required)
# ---------------------------------------------------------------------------


@router.put(
    "/declaration-templates/{provider_type}/{entity_type}",
    response_model=DeclarationTemplateInfo,
)
async def upload_template(
    provider_type: str,
    entity_type: str,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_analyst: Analyst = Depends(require_admin),
    settings: Settings = Depends(get_settings),
):
    """
    Upload or replace the declaration template DOCX for a given provider type
    and entity type. Only DOCX files are accepted. Requires admin privileges.
    """
    if provider_type not in VALID_PROVIDER_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Invalid provider_type. Must be one of: {', '.join(sorted(VALID_PROVIDER_TYPES))}",
        )
    if entity_type not in VALID_ENTITY_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="entity_type must be 'PF' or 'PJ'",
        )

    filename_lower = (file.filename or "").lower()
    content_type_ok = file.content_type == DOCX_MIME_TYPE
    extension_ok = filename_lower.endswith(".docx")
    if not content_type_ok and not extension_ok:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Only DOCX files are accepted for declaration templates",
        )

    # Read content and check size (max 20 MB)
    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="File exceeds the 20 MB size limit",
        )

    # Verify the file is actually a DOCX/ZIP (magic bytes: PK\x03\x04).
    # This prevents someone from uploading a non-DOCX file with a .docx extension.
    if content[:4] != b"PK\x03\x04":
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="File content is not a valid DOCX document",
        )

    # Save to disk — fixed path: declaration_templates/{provider_type}_{entity_type}.docx
    template_dir = _get_template_dir(settings.DOCUMENTS_BASE_PATH)
    file_path = os.path.join(template_dir, f"{provider_type}_{entity_type}.docx")

    safe_original = sanitize_filename(
        file.filename or f"{provider_type}_{entity_type}_declaraciones.docx"
    )

    try:
        with open(file_path, "wb") as f:
            f.write(content)

        # Upsert in DB
        result = await db.execute(
            select(DeclarationTemplate).where(
                DeclarationTemplate.provider_type == provider_type,
                DeclarationTemplate.entity_type == entity_type,
            )
        )
        existing = result.scalar_one_or_none()

        now = datetime.now(timezone.utc)

        if existing:
            existing.file_path = file_path
            existing.original_filename = safe_original
            existing.uploaded_at = now
            existing.uploaded_by_analyst_id = current_analyst.id
            template = existing
        else:
            template = DeclarationTemplate(
                id=uuid.uuid4(),
                provider_type=provider_type,
                entity_type=entity_type,
                file_path=file_path,
                original_filename=safe_original,
                uploaded_at=now,
                uploaded_by_analyst_id=current_analyst.id,
            )
            db.add(template)

        await log_audit(
            db=db,
            action="declaration_template_uploaded",
            analyst_id=current_analyst.id,
            metadata={
                "provider_type": provider_type,
                "entity_type": entity_type,
                "original_filename": safe_original,
                "analyst_email": current_analyst.email,
            },
        )
        await db.commit()
    except Exception:
        # Clean up the file if DB commit failed
        if os.path.exists(file_path):
            os.remove(file_path)
        raise

    await db.refresh(template, ["uploaded_by_analyst"])

    return DeclarationTemplateInfo(
        provider_type=template.provider_type,
        entity_type=template.entity_type,
        provider_type_label=PROVIDER_TYPE_LABELS.get(template.provider_type, template.provider_type),
        entity_type_label=ENTITY_TYPE_LABELS.get(template.entity_type, template.entity_type),
        original_filename=template.original_filename,
        uploaded_at=template.uploaded_at,
        uploaded_by_name=(
            template.uploaded_by_analyst.full_name if template.uploaded_by_analyst else None
        ),
    )
