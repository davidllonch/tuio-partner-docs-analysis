import asyncio
import copy
import io
import json
import logging
import os
import re
import uuid
from datetime import datetime, timezone
from typing import Optional

from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, File, HTTPException, Request, UploadFile, status
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, field_validator
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.auth.jwt import get_current_analyst, require_admin
from app.utils.rate_limit import limiter as _limiter
from app.config import Settings, get_settings
from app.database import get_db
from app.models.analyst import Analyst
from app.models.contract_template import ContractTemplate
from app.utils.audit import log_audit
from app.utils.docx_utils import convert_docx_to_pdf_via_libreoffice
from app.utils.file_utils import sanitize_filename, content_disposition_filename

logger = logging.getLogger(__name__)

router = APIRouter(tags=["contract-templates"])

# Contract templates are only available for these three provider types.
# agencia_seguros is excluded because they use a different contract framework.
VALID_PROVIDER_TYPES = {
    "colaborador_externo",
    "generador_leads",
    "correduria_seguros",
}

VALID_ENTITY_TYPES = {"PF", "PJ"}

PROVIDER_TYPE_LABELS = {
    "correduria_seguros": "Correduría de Seguros",
    "colaborador_externo": "Colaborador Externo",
    "generador_leads": "Generador de Leads",
}

ENTITY_TYPE_LABELS = {
    "PJ": "Persona Jurídica",
    "PF": "Persona Física",
}

# Accepted MIME type for DOCX uploads
DOCX_MIME_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# ---------------------------------------------------------------------------
# Placeholder maps
# ---------------------------------------------------------------------------

# Placeholders filled from partner_info for Persona Física submissions
PARTNER_PF_MAP = {
    "[REPRESENTANTE]": "nombre_apellidos",
    "[NOMBRE Y APELLIDOS DEL PARTNER]": "nombre_apellidos",
    "[PARTNER]": "nombre_apellidos",  # generador_leads PF variant
    "[NIF]": "nif",
    "[DOMICILIO]": "domicilio",
    "[DIRECCIÓN]": "direccion_notificaciones",
    "[EMAIL]": "email",
    "[CONTACTO]": "contacto_notificaciones",
    "[CLAVE DGS]": "clave_dgs",  # correduria only but safe to include for all
    "[CORREDURÍA]": "nombre_apellidos",  # PF contracts: company name = individual's name
    "[CORREDURIA]": "nombre_apellidos",  # non-accented fallback for older template uploads
    "[SOCIEDAD PARTNER]": "nombre_apellidos",  # PF: individual's name as "sociedad"
}

# Placeholders filled from partner_info for Persona Jurídica submissions
PARTNER_PJ_MAP = {
    "[REPRESENTANTE]": "nombre_representante",
    "[NIF]": "nif_representante",
    "[CIF]": "cif",
    "[SOCIEDAD]": "razon_social",
    "[CORREDURÍA]": "razon_social",  # PJ contracts: company name = razón social
    "[CORREDURIA]": "razon_social",  # non-accented fallback for older template uploads
    "[DOMICILIO]": "domicilio_social",
    "[DOMICILIO SOCIAL]": "domicilio_social",
    "[PODER]": "poder",
    "[DIRECCIÓN]": "direccion_notificaciones",
    "[EMAIL]": "email",
    "[CONTACTO]": "contacto_notificaciones",
    "[CLAVE DGS]": "clave_dgs",  # correduria only but safe to include for all
    "[SOCIEDAD PARTNER]": "razon_social",  # PJ: company's razón social
}

# Placeholders filled by the analyst (only used in generate-full)
ANALYST_MAP = {
    "[ACTIVIDAD]": "actividad",
}

# Commission table placeholders — handled separately via row duplication
COMMISSION_PLACEHOLDERS = [
    "[PRODUCTO DE SEGURO]",
    "[PRIMA NETA TRAMO 1]",
    "[COMISIÓN NP]",
    "[COMISIÓN CARTERA]",
]

# Spanish month names for date placeholders
_MESES = [
    "", "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


# ---------------------------------------------------------------------------
# Pydantic schemas
# ---------------------------------------------------------------------------


class ContractTemplateInfo(BaseModel):
    provider_type: str
    entity_type: str
    provider_type_label: str
    entity_type_label: str
    original_filename: str
    uploaded_at: datetime
    uploaded_by_name: Optional[str] = None

    model_config = {"from_attributes": True}


class AllTemplatesResponse(BaseModel):
    templates: list[ContractTemplateInfo]


class GenerateRequest(BaseModel):
    partner_info: dict

    @field_validator("partner_info")
    @classmethod
    def validate_partner_info_values(cls, v):
        if len(v) > 50:
            raise ValueError("partner_info must not contain more than 50 keys")
        for key, val in v.items():
            if isinstance(val, str) and len(val) > 1000:
                raise ValueError(f"Field '{key}' exceeds maximum length of 1000 characters")
        return v


class GenerateFullRequest(BaseModel):
    partner_info: dict
    contract_data: dict  # { "fields": { "actividad": "..." }, "commissions": [...] }

    # S12: prevent any single field value from being excessively long
    @field_validator("partner_info")
    @classmethod
    def validate_partner_info_values(cls, v):
        if len(v) > 50:
            raise ValueError("partner_info must not contain more than 50 keys")
        for key, val in v.items():
            if isinstance(val, str) and len(val) > 1000:
                raise ValueError(f"Field '{key}' exceeds maximum length of 1000 characters")
        return v

    @field_validator("contract_data")
    @classmethod
    def validate_contract_data_size(cls, v):
        if len(json.dumps(v)) > 100_000:
            raise ValueError("contract_data exceeds maximum allowed size")
        commissions = v.get("commissions", [])
        if len(commissions) > 100:
            raise ValueError("commissions must not exceed 100 rows")
        fields = v.get("fields", {})
        if len(fields) > 50:
            raise ValueError("contract_data.fields must not contain more than 50 keys")
        return v


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _write_bytes_to_file(path: str, data: bytes) -> None:
    """Synchronous file write — called via asyncio.to_thread to avoid blocking."""
    with open(path, "wb") as f:
        f.write(data)


async def _get_template_dir(documents_base_path: str) -> str:
    template_dir = os.path.join(documents_base_path, "contract_templates")
    await asyncio.to_thread(os.makedirs, template_dir, exist_ok=True)
    return template_dir


# OOXML Word namespace — used to find <w:t> elements via lxml
_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# ---------------------------------------------------------------------------
# Encoding-tolerant helpers
# ---------------------------------------------------------------------------
# Some DOCX files store accented characters (e.g. Í, Ó, Á) as single Latin-1
# bytes instead of proper 2-byte UTF-8 sequences. Because the XML header still
# declares UTF-8, lxml replaces those invalid bytes with U+FFFD (the Unicode
# replacement character, '?'). This means our placeholder strings (which carry
# the correct Unicode characters) won't match the lxml-read text.
#
# _fuzzy_eq and _fuzzy_pattern handle this by treating U+FFFD as a wildcard
# that matches any single character, so comparisons and regex substitutions
# work regardless of the encoding corruption.
_REPL = "\ufffd"  # U+FFFD — Unicode replacement character


def _fuzzy_eq(a: str, b: str) -> bool:
    """True if strings are equal treating U+FFFD in either string as wildcard."""
    if len(a) != len(b):
        return False
    return all(ca == cb or ca == _REPL or cb == _REPL for ca, cb in zip(a, b))


def _fuzzy_sub(pattern_text: str, replacement: str, subject: str) -> str:
    """
    Replace the first (or all) occurrences of pattern_text in subject,
    where non-ASCII chars in pattern_text also match U+FFFD in subject.
    Falls back to plain str.replace when no non-ASCII chars are present
    (avoids regex overhead for ASCII-only placeholders like [DIA]).
    """
    # Fast path: no non-ASCII characters in the placeholder
    if all(ord(ch) < 128 for ch in pattern_text):
        return subject.replace(pattern_text, replacement)

    # Build regex where each non-ASCII char also accepts U+FFFD
    parts = []
    for ch in pattern_text:
        if ch == _REPL:
            parts.append(".")
        elif ord(ch) > 127:
            parts.append(f"(?:{re.escape(ch)}|{re.escape(_REPL)})")
        else:
            parts.append(re.escape(ch))
    pat = re.compile("".join(parts))
    return pat.sub(replacement, subject)


def _fuzzy_search(pattern_text: str, subject: str) -> bool:
    """Return True if subject contains pattern_text (fuzzy-encoded)."""
    # Fast path
    if all(ord(ch) < 128 for ch in pattern_text):
        return pattern_text in subject
    # Check both the literal and the \ufffd-substituted variant
    if pattern_text in subject:
        return True
    # Build fuzzy regex
    parts = []
    for ch in pattern_text:
        if ch == _REPL:
            parts.append(".")
        elif ord(ch) > 127:
            parts.append(f"(?:{re.escape(ch)}|{re.escape(_REPL)})")
        else:
            parts.append(re.escape(ch))
    return bool(re.search("".join(parts), subject))


def _replace_in_paragraph_elem(p_elem, replacements: dict[str, str]) -> None:
    """
    Replace placeholder strings in a single paragraph XML element.

    Three-pass strategy:
      Pass 1 — replace within each individual <w:t> element (handles placeholders
               fully contained in one run; preserves per-run formatting).
      Pass 2 — sliding-window for split-run placeholders: handles the common Word
               pattern where the opening '[' is the last char of run N, the inner
               name is the entirety of run N+1, and ']' is the first char of run N+2.
               Only the three involved characters are modified; all other runs keep
               their original text and formatting unchanged.
      Pass 3 — full-collapse fallback for any remaining cross-run placeholders.
               Loses per-run formatting for the paragraph but always works.

    All passes use fuzzy matching so that U+FFFD (the Unicode replacement character
    produced by lxml when a DOCX contains invalid UTF-8 bytes for accented chars)
    is treated as a wildcard matching the expected character.
    """
    w_t_elems = p_elem.findall(".//" + _W_NS + "t")
    if not w_t_elems:
        return

    # Pass 1: per-run replacement (preserves formatting)
    for elem in w_t_elems:
        if not elem.text:
            continue
        for placeholder, value in replacements.items():
            if _fuzzy_search(placeholder, elem.text):
                elem.text = _fuzzy_sub(placeholder, value, elem.text)

    # Pass 2: sliding-window for split-run placeholders
    # Pattern: w_t[i-1] ends with "[", w_t[i] ≈ "INNER_NAME", w_t[i+1] starts with "]"
    for placeholder, value in replacements.items():
        if len(placeholder) < 3 or placeholder[0] != "[" or placeholder[-1] != "]":
            continue
        inner = placeholder[1:-1]  # strip the [ and ]
        i = 1
        while i < len(w_t_elems) - 1:
            curr_text = w_t_elems[i].text or ""
            if _fuzzy_eq(curr_text, inner):  # fuzzy: Í matches \ufffd
                prev_text = w_t_elems[i - 1].text or ""
                next_text = w_t_elems[i + 1].text or ""
                if prev_text.endswith("[") and next_text.startswith("]"):
                    # Found a split placeholder — fix it in-place
                    w_t_elems[i - 1].text = prev_text[:-1]   # strip trailing [
                    w_t_elems[i].text = value                # fill replacement
                    w_t_elems[i + 1].text = next_text[1:]    # strip leading ]
                    i += 2  # skip past the replaced placeholder
                    continue
            i += 1

    # Pass 3: full-collapse fallback for any remaining cross-run placeholders
    full_text = "".join(elem.text or "" for elem in w_t_elems)
    remaining = {ph: v for ph, v in replacements.items() if _fuzzy_search(ph, full_text)}
    if not remaining:
        return

    new_text = full_text
    for placeholder, value in remaining.items():
        new_text = _fuzzy_sub(placeholder, value, new_text)
    w_t_elems[0].text = new_text
    for elem in w_t_elems[1:]:
        elem.text = ""


def _iter_all_tables_inline(doc_or_cells):
    """
    BFS generator that yields every table in the document including nested ones.
    Accepts either a Document (uses doc.tables) or a list of cells.
    This is defined early so _replace_placeholders_in_docx can use it.
    """
    if hasattr(doc_or_cells, "tables"):
        queue = list(doc_or_cells.tables)
    else:
        queue = list(doc_or_cells)
    while queue:
        table = queue.pop(0)
        yield table
        for row in table.rows:
            for cell in row.cells:
                queue.extend(cell.tables)


def _iter_all_tables_inline_from_tables(tables):
    """Same as _iter_all_tables_inline but accepts a plain list of tables."""
    queue = list(tables)
    while queue:
        table = queue.pop(0)
        yield table
        for row in table.rows:
            for cell in row.cells:
                queue.extend(cell.tables)


def _strip_all_highlights(doc: DocxDocument) -> None:
    """
    Remove every <w:shd> and <w:highlight> element from the entire document body,
    including all headers and footers.

    DOCX templates use cell shading, paragraph shading, run shading, or run
    highlighting (yellow/orange/etc.) to mark placeholder fields visually. Once
    all placeholders are replaced this function cleans the document so the final
    PDF has no leftover colour backgrounds.

    Covers all three levels where colour can live:
      • <w:tcPr><w:shd>   — table cell background
      • <w:pPr><w:shd>    — paragraph background
      • <w:rPr><w:shd>    — run background shading
      • <w:rPr><w:highlight> — run highlight (yellow, green, etc.)
    """
    _STRIP_TAGS = {_W_NS + "shd", _W_NS + "highlight"}

    def _strip_tree(root_elem) -> None:
        # Collect first to avoid mutating while iterating
        to_remove = [e for e in root_elem.iter() if e.tag in _STRIP_TAGS]
        for elem in to_remove:
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)

    # Main document body
    _strip_tree(doc.element.body)

    # Headers and footers for every section
    for section in doc.sections:
        for hdr_ftr in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            try:
                if hdr_ftr is not None and hdr_ftr._element is not None:
                    _strip_tree(hdr_ftr._element)
            except Exception:
                pass  # Some sections may not have all header/footer variants


def _strip_highlights_from_filled_elements(doc: DocxDocument) -> None:
    """
    Remove highlight colours ONLY from elements whose text no longer contains
    an unfilled placeholder marker (i.e. a '[' character).

    Used for the partner PDF download: placeholders that the partner has
    already filled in via the form are shown without colour (the value is
    readable as normal text), while placeholders that are still pending
    (commission rows, [ACTIVIDAD], [SÍ/NO], etc.) keep their yellow/cyan
    highlight so the partner knows these will be completed later.

    Granularity:
      - Run level: strip <w:rPr><w:shd> and <w:rPr><w:highlight> from any
        run whose text contains no '['.
      - Cell level: strip <w:tcPr><w:shd> from any table cell whose entire
        subtree contains no '['.
      - Paragraph level: strip <w:pPr><w:shd> from any paragraph whose text
        contains no '['.
    """
    _STRIP_TAGS = {_W_NS + "shd", _W_NS + "highlight"}

    def _elem_has_placeholder(elem) -> bool:
        """True if any <w:t> descendant of elem still contains a '[' char."""
        return any(
            t.text and "[" in t.text
            for t in elem.iter(_W_NS + "t")
        )

    def _strip_from_elem(elem) -> None:
        to_remove = [e for e in elem.iter() if e.tag in _STRIP_TAGS]
        for e in to_remove:
            parent = e.getparent()
            if parent is not None:
                parent.remove(e)

    def _process(root_elem) -> None:
        # Strip cell-level shading for cells with no remaining placeholders
        for tc in root_elem.iter(_W_NS + "tc"):
            if not _elem_has_placeholder(tc):
                tc_pr = tc.find(_W_NS + "tcPr")
                if tc_pr is not None:
                    shd = tc_pr.find(_W_NS + "shd")
                    if shd is not None:
                        tc_pr.remove(shd)

        # Strip paragraph-level shading for paragraphs with no remaining placeholders
        for p in root_elem.iter(_W_NS + "p"):
            if not _elem_has_placeholder(p):
                p_pr = p.find(_W_NS + "pPr")
                if p_pr is not None:
                    shd = p_pr.find(_W_NS + "shd")
                    if shd is not None:
                        p_pr.remove(shd)

        # Strip run-level shading/highlight for runs in paragraphs with no remaining placeholders
        for p in root_elem.iter(_W_NS + "p"):
            if _elem_has_placeholder(p):
                continue  # paragraph still has unfilled placeholder(s) — keep all run colours
            for r in p.iter(_W_NS + "r"):
                r_pr = r.find(_W_NS + "rPr")
                if r_pr is not None:
                    for tag in (_W_NS + "shd", _W_NS + "highlight"):
                        e = r_pr.find(tag)
                        if e is not None:
                            r_pr.remove(e)

    _process(doc.element.body)
    for section in doc.sections:
        for hdr_ftr in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            try:
                if hdr_ftr is not None and hdr_ftr._element is not None:
                    _process(hdr_ftr._element)
            except Exception:
                pass


def _replace_placeholders_in_docx(doc: DocxDocument, replacements: dict[str, str]) -> None:
    """
    Replace placeholder strings (e.g. '[CAMPO]') everywhere in a python-docx Document.

    Uses lxml .iter() to traverse ALL <w:p> elements regardless of nesting level.
    This correctly handles placeholders inside:
      - Body paragraphs
      - Table cells (including nested tables)
      - Structured Document Tags (<w:sdt>) — Word content controls / locked regions.
        python-docx's doc.paragraphs and doc.tables skip SDT content, so lxml
        traversal is the only way to reach placeholders like [CORREDURÍA] that
        Word places inside SDTs in signature areas.
      - Document headers and footers for every section
    """
    # Main body — iterate every <w:p> element, including those inside SDTs and nested tables
    for p_elem in doc.element.body.iter(_W_NS + "p"):
        _replace_in_paragraph_elem(p_elem, replacements)

    # Headers and footers
    for section in doc.sections:
        for hdr_ftr in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            try:
                if hdr_ftr is not None and hdr_ftr._element is not None:
                    for p_elem in hdr_ftr._element.iter(_W_NS + "p"):
                        _replace_in_paragraph_elem(p_elem, replacements)
            except Exception:
                pass


def _build_partner_replacements(entity_type: str, partner_info: dict) -> dict[str, str]:
    """Build placeholder → value mapping for partner fields + date."""
    today = datetime.now(timezone.utc)
    date_replacements = {
        "[DÍA]": str(today.day),
        "[MES]": _MESES[today.month],
        "[AÑO]": str(today.year),
    }

    field_map = PARTNER_PF_MAP if entity_type == "PF" else PARTNER_PJ_MAP

    replacements: dict[str, str] = {}
    for placeholder, field_key in field_map.items():
        value = partner_info.get(field_key, "")
        replacements[placeholder] = str(value)

    replacements.update(date_replacements)
    return replacements


def _build_full_replacements(
    entity_type: str,
    partner_info: dict,
    contract_fields: dict,
) -> dict[str, str]:
    """Build placeholder → value mapping for partner fields + analyst fields + date."""
    replacements = _build_partner_replacements(entity_type, partner_info)

    for placeholder, field_key in ANALYST_MAP.items():
        value = contract_fields.get(field_key, "")
        replacements[placeholder] = str(value)

    return replacements


# ---------------------------------------------------------------------------
# Commission row duplication helpers
# ---------------------------------------------------------------------------


def _replace_commission_placeholders_in_row(row, commission: dict) -> None:
    """Replace commission placeholders in an existing python-docx table row."""
    replacements = {
        "[PRODUCTO DE SEGURO]": commission.get("producto", ""),
        "[PRIMA NETA TRAMO 1]": commission.get("prima", ""),
        "[COMISIÓN NP]": commission.get("comision_np", ""),
        "[COMISIÓN CARTERA]": commission.get("comision_cartera", ""),
    }
    for cell in row.cells:
        for para in cell.paragraphs:
            _replace_in_paragraph_elem(para._p, replacements)


def _replace_commission_in_tr(tr_elem, commission: dict) -> None:
    """Replace commission placeholders directly in a raw lxml <w:tr> element."""
    replacements = {
        "[PRODUCTO DE SEGURO]": commission.get("producto", ""),
        "[PRIMA NETA TRAMO 1]": commission.get("prima", ""),
        "[COMISIÓN NP]": commission.get("comision_np", ""),
        "[COMISIÓN CARTERA]": commission.get("comision_cartera", ""),
    }
    # Use _replace_in_paragraph_elem for each <w:p> inside this row,
    # so split-run placeholders are also handled correctly.
    for p_elem in tr_elem.findall(".//" + _W_NS + "p"):
        _replace_in_paragraph_elem(p_elem, replacements)
    # Remove background highlight from all cells in this row
    for tc in tr_elem.findall(".//" + _W_NS + "tc"):
        tc_pr = tc.find(_W_NS + "tcPr")
        if tc_pr is not None:
            shd = tc_pr.find(_W_NS + "shd")
            if shd is not None:
                tc_pr.remove(shd)


def _iter_all_tables(doc: DocxDocument):
    """
    Yield all tables in the document, including nested ones inside table cells.
    python-docx's doc.tables only returns top-level tables; nested tables require
    recursive traversal via cell.tables.
    """
    queue = list(doc.tables)
    while queue:
        table = queue.pop(0)
        yield table
        for row in table.rows:
            for cell in row.cells:
                queue.extend(cell.tables)


def _fill_commission_rows(doc: DocxDocument, commission_rows: list[dict]) -> None:
    """
    Find the Tramo 1 commission template row, fill it with commission data, and
    delete any adjacent Tramo 2 template row (the DOCX has two template rows per
    product, but we only use Tramo 1 — Tramo 2 is always removed).

    For N commission entries:
      - The first entry fills the original template row in-place.
      - Entries 2..N get a deep-copied duplicate inserted below.
      - The Tramo 2 row (if present immediately after the Tramo 1 row) is always
        deleted regardless of how many commissions are provided.
      - If commission_rows is empty, the Tramo 1 template row is also deleted.

    Uses _iter_all_tables to also search inside nested tables.
    """
    for table in _iter_all_tables(doc):
        for i, row in enumerate(table.rows):
            row_text = "".join(cell.text for cell in row.cells)
            if "[PRODUCTO DE SEGURO]" not in row_text:
                continue

            template_tr = row._tr
            parent = template_tr.getparent()

            # Detect and remove the Tramo 2 row (immediately following, if present)
            tramo2_tr = None
            if i + 1 < len(table.rows):
                next_row = table.rows[i + 1]
                next_text = "".join(cell.text for cell in next_row.cells)
                if "[PRIMA NETA TRAMO 2]" in next_text or (
                    "[PRODUCTO DE SEGURO]" in next_text and next_row is not row
                ):
                    tramo2_tr = next_row._tr

            if not commission_rows:
                # No data — remove template row (and Tramo 2 if present)
                parent.remove(template_tr)
                if tramo2_tr is not None:
                    parent.remove(tramo2_tr)
                return

            # ① Deep copy the Tramo 1 template row BEFORE modifying it.
            #    One copy per additional commission (commissions[1:]).
            insert_after = template_tr
            additional = []
            for commission in commission_rows[1:]:
                additional.append((copy.deepcopy(template_tr), commission))

            # ② Fill the original Tramo 1 template row with the first commission.
            _replace_commission_in_tr(template_tr, commission_rows[0])

            # ③ Insert and fill copies for remaining commissions.
            for new_tr, commission in additional:
                insert_idx = list(parent).index(insert_after) + 1
                parent.insert(insert_idx, new_tr)
                insert_after = new_tr
                _replace_commission_in_tr(new_tr, commission)

            # ④ Remove the Tramo 2 row (always, regardless of commissions).
            if tramo2_tr is not None:
                parent.remove(tramo2_tr)

            return  # Only process the first matching commission table


def _extract_placeholder_context(doc: DocxDocument, placeholder: str) -> str | None:
    """
    Return the full text of the first paragraph that contains the given placeholder.
    Searches both body paragraphs and table cells (including nested tables).
    Returns None if the placeholder is not found.
    """
    for para in doc.paragraphs:
        if placeholder in para.text:
            return para.text.strip()
    for table in _iter_all_tables(doc):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        return para.text.strip()
    return None


def _extract_si_no_fields(doc: DocxDocument) -> list[str]:
    """
    Return the label of every table row (or paragraph) that contains a SI/NO
    placeholder.  Accepts all variants: "[SI/NO]", "[SI / NO]", "[SÍ/NO]",
    "[SÍ / NO]".

    Uses lxml .iter() so it reaches paragraphs inside Structured Document Tags
    (<w:sdt>) and deeply nested tables that python-docx's doc.tables / doc.paragraphs
    do not expose.  For each matching paragraph:
      - If it lives inside a <w:tc> we navigate up to the <w:tr> and use the
        first cell's text as the product label.
      - Otherwise we use the text before the first "[" in the paragraph.
    """
    _SI_NO_VARIANTS = ["[SI/NO]", "[SI / NO]", "[SÍ/NO]", "[SÍ / NO]"]
    results: list[str] = []

    for p_elem in doc.element.body.iter(_W_NS + "p"):
        full_text = "".join(t.text or "" for t in p_elem.iter(_W_NS + "t"))
        if not any(_fuzzy_search(v, full_text) for v in _SI_NO_VARIANTS):
            continue

        # Try to get the label from the first cell of the enclosing table row
        label = None
        ancestor = p_elem.getparent()
        while ancestor is not None:
            if ancestor.tag == _W_NS + "tr":
                first_tc = ancestor.find(_W_NS + "tc")
                if first_tc is not None:
                    label = "".join(
                        t.text or "" for t in first_tc.iter(_W_NS + "t")
                    ).strip()
                break
            ancestor = ancestor.getparent()

        # Fallback: text before the first "[" in the paragraph
        if not label:
            bracket_pos = full_text.find("[")
            if bracket_pos > 0:
                label = full_text[:bracket_pos].strip().rstrip(":").strip()
            else:
                label = full_text.strip()

        if label and label not in results:
            results.append(label)

    return results


def _fill_si_no_fields(doc: DocxDocument, si_no_values: dict[str, str]) -> None:
    """
    For each paragraph containing a SI/NO placeholder, look up the product
    label (from the first cell of the enclosing table row, or from text before
    the "[") in si_no_values and replace [SÍ / NO] with "Sí" or "No".

    Uses lxml .iter() so it reaches paragraphs inside <w:sdt> elements and
    deeply nested tables that python-docx's normal traversal misses.
    """
    _SI_NO_VARIANTS = ["[SI/NO]", "[SI / NO]", "[SÍ/NO]", "[SÍ / NO]"]
    replacements_base = {v: "" for v in _SI_NO_VARIANTS}  # filled per paragraph

    for p_elem in doc.element.body.iter(_W_NS + "p"):
        full_text = "".join(t.text or "" for t in p_elem.iter(_W_NS + "t"))
        if not any(_fuzzy_search(v, full_text) for v in _SI_NO_VARIANTS):
            continue

        # Determine label
        label = None
        ancestor = p_elem.getparent()
        while ancestor is not None:
            if ancestor.tag == _W_NS + "tr":
                first_tc = ancestor.find(_W_NS + "tc")
                if first_tc is not None:
                    label = "".join(
                        t.text or "" for t in first_tc.iter(_W_NS + "t")
                    ).strip()
                break
            ancestor = ancestor.getparent()

        if not label:
            bracket_pos = full_text.find("[")
            if bracket_pos > 0:
                label = full_text[:bracket_pos].strip().rstrip(":").strip()
            else:
                label = full_text.strip()

        value = si_no_values.get(label or "", "")
        if value:
            replacements = {v: value for v in _SI_NO_VARIANTS}
            _replace_in_paragraph_elem(p_elem, replacements)


# ---------------------------------------------------------------------------
# GET /api/contract-templates  (JWT required — list all for admin)
# ---------------------------------------------------------------------------


@router.get("/contract-templates", response_model=AllTemplatesResponse)
async def list_all_templates(
    db: AsyncSession = Depends(get_db),
    current_analyst: Analyst = Depends(get_current_analyst),
):
    """
    Return all contract templates (one per provider_type × entity_type combination).
    Used by the admin UI to show what has been uploaded.
    """
    result = await db.execute(
        select(ContractTemplate).options(
            selectinload(ContractTemplate.uploaded_by_analyst)
        )
    )
    templates = result.scalars().all()

    template_map = {
        (t.provider_type, t.entity_type): t for t in templates
    }

    items = []
    for pt in sorted(VALID_PROVIDER_TYPES):
        for et in ("PJ", "PF"):
            t = template_map.get((pt, et))
            if t:
                items.append(
                    ContractTemplateInfo(
                        provider_type=t.provider_type,
                        entity_type=t.entity_type,
                        provider_type_label=PROVIDER_TYPE_LABELS.get(t.provider_type, t.provider_type),
                        entity_type_label=ENTITY_TYPE_LABELS.get(t.entity_type, t.entity_type),
                        original_filename=t.original_filename,
                        uploaded_at=t.uploaded_at,
                        uploaded_by_name=(
                            t.uploaded_by_analyst.full_name
                            if t.uploaded_by_analyst
                            else None
                        ),
                    )
                )

    return AllTemplatesResponse(templates=items)


# ---------------------------------------------------------------------------
# GET /api/contract-templates/{provider_type}/{entity_type}  (PUBLIC)
# ---------------------------------------------------------------------------


@router.get("/contract-templates/{provider_type}/{entity_type}")
async def get_template_info(
    provider_type: str,
    entity_type: str,
    db: AsyncSession = Depends(get_db),
):
    """
    Public endpoint: returns metadata about a contract template if it exists.
    Used by the partner invite page to decide whether to show a download button.
    Returns 404 if no template has been uploaded for this combination.
    """
    if provider_type not in VALID_PROVIDER_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Invalid provider_type. Must be one of: {', '.join(sorted(VALID_PROVIDER_TYPES))}",
        )
    if entity_type not in VALID_ENTITY_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="entity_type must be 'PF' or 'PJ'",
        )

    result = await db.execute(
        select(ContractTemplate).where(
            ContractTemplate.provider_type == provider_type,
            ContractTemplate.entity_type == entity_type,
        )
    )
    template = result.scalar_one_or_none()

    if template is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="No template uploaded")

    return {
        "provider_type": template.provider_type,
        "entity_type": template.entity_type,
        "original_filename": template.original_filename,
        "uploaded_at": template.uploaded_at.isoformat(),
    }


# ---------------------------------------------------------------------------
# GET /api/contract-templates/{provider_type}/{entity_type}/download  (PUBLIC)
# ---------------------------------------------------------------------------


@router.get("/contract-templates/{provider_type}/{entity_type}/download")
@_limiter.limit("60/hour")
async def download_template(
    request: Request,
    provider_type: str,
    entity_type: str,
    db: AsyncSession = Depends(get_db),
    settings: Settings = Depends(get_settings),
):
    """
    Public endpoint: download the raw DOCX template for a provider type + entity type.
    """
    if provider_type not in VALID_PROVIDER_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Invalid provider_type",
        )
    if entity_type not in VALID_ENTITY_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="entity_type must be 'PF' or 'PJ'",
        )

    result = await db.execute(
        select(ContractTemplate).where(
            ContractTemplate.provider_type == provider_type,
            ContractTemplate.entity_type == entity_type,
        )
    )
    template = result.scalar_one_or_none()

    if template is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="No template found")

    # Prevent path traversal: ensure the stored path resolves inside the expected dir
    template_dir = os.path.join(settings.DOCUMENTS_BASE_PATH, "contract_templates")
    real_path = os.path.realpath(template.file_path)
    if not real_path.startswith(os.path.realpath(template_dir) + os.sep):
        logger.error("Contract template path escapes base dir: %s", template.file_path)
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Template not found")

    if not await asyncio.to_thread(os.path.exists, real_path):
        logger.error(
            "Contract template file missing on disk: %s (provider=%s entity=%s)",
            template.file_path,
            provider_type,
            entity_type,
        )
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Template file not found on disk",
        )

    return FileResponse(
        path=real_path,
        media_type=DOCX_MIME_TYPE,
        filename=template.original_filename,
        headers={
            "Content-Disposition": content_disposition_filename(template.original_filename),
        },
    )


# ---------------------------------------------------------------------------
# POST /api/contract-templates/{provider_type}/{entity_type}/generate  (PUBLIC)
# ---------------------------------------------------------------------------


@router.post("/contract-templates/{provider_type}/{entity_type}/generate")
@_limiter.limit("20/hour")
async def generate_contract_pdf(
    request: Request,
    provider_type: str,
    entity_type: str,
    body: GenerateRequest,
    db: AsyncSession = Depends(get_db),
    settings: Settings = Depends(get_settings),
):
    """
    Public endpoint (partner version): fill partner fields + date in the contract
    template and return a personalised PDF.

    Commission placeholders ([PRODUCTO DE SEGURO], etc.) and analyst fields
    ([ACTIVIDAD]) are intentionally left unchanged — they will be filled later
    by the analyst via the /generate-full endpoint.

    Pipeline:
      1. Load DOCX from disk
      2. Replace partner placeholders only (lxml approach)
      3. Convert patched DOCX → PDF via LibreOffice headless subprocess
      4. Stream the PDF back to the caller
    """
    if provider_type not in VALID_PROVIDER_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Invalid provider_type. Must be one of: {', '.join(sorted(VALID_PROVIDER_TYPES))}",
        )
    if entity_type not in VALID_ENTITY_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="entity_type must be 'PF' or 'PJ'",
        )

    result = await db.execute(
        select(ContractTemplate).where(
            ContractTemplate.provider_type == provider_type,
            ContractTemplate.entity_type == entity_type,
        )
    )
    template = result.scalar_one_or_none()

    if template is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No contract template found for this provider type and entity type",
        )

    # Prevent path traversal: ensure the stored path resolves inside the expected dir
    generate_template_dir = os.path.join(settings.DOCUMENTS_BASE_PATH, "contract_templates")
    real_generate_path = os.path.realpath(template.file_path)
    if not real_generate_path.startswith(os.path.realpath(generate_template_dir) + os.sep):
        logger.error("Contract template path escapes base dir: %s", template.file_path)
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Template not found")

    if not await asyncio.to_thread(os.path.exists, real_generate_path):
        logger.error(
            "Contract template file missing on disk: %s", template.file_path
        )
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Template file not found on disk",
        )

    # ── 1. Load and patch the DOCX (partner fields + date only) ──────────────
    # DocxDocument opens a ZIP archive and parses XML — offload to thread pool.
    doc = await asyncio.to_thread(DocxDocument, real_generate_path)
    replacements = _build_partner_replacements(entity_type, body.partner_info)
    _replace_placeholders_in_docx(doc, replacements)
    # Commission placeholders and [ACTIVIDAD] are deliberately left as-is.
    # Strip highlights only from fields that have been filled; keep colours on
    # any remaining '[PLACEHOLDER]' text so the partner can see what's pending.
    _strip_highlights_from_filled_elements(doc)

    # ── 2. Save patched DOCX to an in-memory buffer ───────────────────────────
    # doc.save() serialises XML + ZIP — offload to thread pool.
    docx_buffer = io.BytesIO()
    await asyncio.to_thread(doc.save, docx_buffer)
    docx_bytes = docx_buffer.getvalue()

    # ── 3. Convert patched DOCX → PDF via LibreOffice ────────────────────────
    pdf_bytes = await convert_docx_to_pdf_via_libreoffice(docx_bytes)

    # ── 4. Stream back ────────────────────────────────────────────────────────
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={
            "Content-Disposition": 'inline; filename="contrato.pdf"',
        },
    )


# ---------------------------------------------------------------------------
# POST /api/contract-templates/{provider_type}/{entity_type}/generate-full  (JWT)
# ---------------------------------------------------------------------------


@router.post("/contract-templates/{provider_type}/{entity_type}/generate-full")
@_limiter.limit("30/hour")
async def generate_full_contract_pdf(
    request: Request,
    provider_type: str,
    entity_type: str,
    body: GenerateFullRequest,
    db: AsyncSession = Depends(get_db),
    settings: Settings = Depends(get_settings),
    current_analyst: Analyst = Depends(get_current_analyst),
):
    """
    Analyst-only endpoint: fill all placeholders (partner + analyst fields + commissions)
    and return a fully completed contract PDF.

    Pipeline:
      1. Load DOCX from disk
      2. Replace all partner + analyst placeholders
      3. Duplicate commission table rows and fill them
      4. Convert to PDF via LibreOffice
      5. Stream the PDF back
    """
    if provider_type not in VALID_PROVIDER_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Invalid provider_type. Must be one of: {', '.join(sorted(VALID_PROVIDER_TYPES))}",
        )
    if entity_type not in VALID_ENTITY_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="entity_type must be 'PF' or 'PJ'",
        )

    result = await db.execute(
        select(ContractTemplate).where(
            ContractTemplate.provider_type == provider_type,
            ContractTemplate.entity_type == entity_type,
        )
    )
    template = result.scalar_one_or_none()

    if template is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No contract template found for this provider type and entity type",
        )

    # Prevent path traversal: ensure the stored path resolves inside the expected dir
    full_template_dir = os.path.join(settings.DOCUMENTS_BASE_PATH, "contract_templates")
    real_full_path = os.path.realpath(template.file_path)
    if not real_full_path.startswith(os.path.realpath(full_template_dir) + os.sep):
        logger.error("Contract template path escapes base dir: %s", template.file_path)
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Template not found")

    if not await asyncio.to_thread(os.path.exists, real_full_path):
        logger.error(
            "Contract template file missing on disk: %s", template.file_path
        )
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Template file not found on disk",
        )

    contract_fields = body.contract_data.get("fields", {})
    commission_rows = body.contract_data.get("commissions", [])

    # ── 1. Load DOCX ─────────────────────────────────────────────────────────
    # DocxDocument opens a ZIP archive and parses XML — offload to thread pool.
    doc = await asyncio.to_thread(DocxDocument, real_full_path)

    # ── 2. Replace all partner + analyst placeholders ─────────────────────────
    replacements = _build_full_replacements(entity_type, body.partner_info, contract_fields)
    _replace_placeholders_in_docx(doc, replacements)

    # ── 3. Handle SI/NO fields (Annex I) ─────────────────────────────────────
    si_no_fields = body.contract_data.get("si_no_fields", {})
    if si_no_fields:
        _fill_si_no_fields(doc, si_no_fields)

    # ── 4. Handle commission table rows ──────────────────────────────────────
    # Filter out empty rows (producto is required; rows with no producto are skipped)
    non_empty_commissions = [
        r for r in commission_rows if r.get("producto", "").strip()
    ]
    # Always call _fill_commission_rows — if the list is empty it removes the template row
    _fill_commission_rows(doc, non_empty_commissions)

    # ── 5. Strip all template highlight colours ───────────────────────────────
    _strip_all_highlights(doc)

    # ── 6. Save patched DOCX to an in-memory buffer ───────────────────────────
    # doc.save() serialises XML + ZIP — offload to thread pool.
    docx_buffer = io.BytesIO()
    await asyncio.to_thread(doc.save, docx_buffer)
    docx_bytes = docx_buffer.getvalue()

    # ── 7. Convert patched DOCX → PDF via LibreOffice ────────────────────────
    pdf_bytes = await convert_docx_to_pdf_via_libreoffice(docx_bytes)

    # ── 8. Stream back ────────────────────────────────────────────────────────
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={
            "Content-Disposition": 'inline; filename="contrato_completo.pdf"',
        },
    )


# ---------------------------------------------------------------------------
# GET /api/contract-templates/{provider_type}/{entity_type}/placeholder-context  (JWT)
# ---------------------------------------------------------------------------


@router.get("/contract-templates/{provider_type}/{entity_type}/placeholder-context")
async def get_placeholder_context(
    provider_type: str,
    entity_type: str,
    db: AsyncSession = Depends(get_db),
    settings: Settings = Depends(get_settings),
    current_analyst: Analyst = Depends(get_current_analyst),
):
    """
    Return the full paragraph text surrounding key placeholders ([ACTIVIDAD]).
    Used by the analyst UI to show context hints next to input fields.
    Returns an empty dict if no template has been uploaded.
    """
    if provider_type not in VALID_PROVIDER_TYPES:
        raise HTTPException(status_code=422, detail="Invalid provider_type")
    if entity_type not in VALID_ENTITY_TYPES:
        raise HTTPException(status_code=422, detail="entity_type must be 'PF' or 'PJ'")

    result = await db.execute(
        select(ContractTemplate).where(
            ContractTemplate.provider_type == provider_type,
            ContractTemplate.entity_type == entity_type,
        )
    )
    template = result.scalar_one_or_none()
    if template is None or not await asyncio.to_thread(os.path.exists, template.file_path):
        return {"context": {}}

    try:
        doc = DocxDocument(template.file_path)
        context: dict[str, str] = {}
        actividad_ctx = _extract_placeholder_context(doc, "[ACTIVIDAD]")
        if actividad_ctx:
            context["ACTIVIDAD"] = actividad_ctx
        return {"context": context}
    except Exception as exc:
        logger.error("Error extracting placeholder context: %s", exc)
        return {"context": {}}


# ---------------------------------------------------------------------------
# GET /api/contract-templates/{provider_type}/{entity_type}/si-no-fields  (JWT)
# ---------------------------------------------------------------------------


@router.get("/contract-templates/{provider_type}/{entity_type}/si-no-fields")
async def get_si_no_fields(
    provider_type: str,
    entity_type: str,
    db: AsyncSession = Depends(get_db),
    settings: Settings = Depends(get_settings),
    current_analyst: Analyst = Depends(get_current_analyst),
):
    """
    Return the list of insurance product labels that have a [SI/NO] placeholder
    in the contract template's Annex I table.
    Returns an empty list if no template has been uploaded.
    """
    if provider_type not in VALID_PROVIDER_TYPES:
        raise HTTPException(status_code=422, detail="Invalid provider_type")
    if entity_type not in VALID_ENTITY_TYPES:
        raise HTTPException(status_code=422, detail="entity_type must be 'PF' or 'PJ'")

    result = await db.execute(
        select(ContractTemplate).where(
            ContractTemplate.provider_type == provider_type,
            ContractTemplate.entity_type == entity_type,
        )
    )
    template = result.scalar_one_or_none()
    if template is None or not await asyncio.to_thread(os.path.exists, template.file_path):
        return {"fields": []}

    try:
        doc = DocxDocument(template.file_path)
        fields = _extract_si_no_fields(doc)
        return {"fields": fields}
    except Exception as exc:
        logger.error("Error extracting SI/NO fields: %s", exc)
        return {"fields": []}


# ---------------------------------------------------------------------------
# PUT /api/contract-templates/{provider_type}/{entity_type}  (JWT required)
# ---------------------------------------------------------------------------


@router.put(
    "/contract-templates/{provider_type}/{entity_type}",
    response_model=ContractTemplateInfo,
)
async def upload_template(
    provider_type: str,
    entity_type: str,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_analyst: Analyst = Depends(require_admin),
    settings: Settings = Depends(get_settings),
):
    """
    Upload or replace the contract template DOCX for a given provider type
    and entity type. Only DOCX files are accepted. Requires admin privileges.
    """
    if provider_type not in VALID_PROVIDER_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Invalid provider_type. Must be one of: {', '.join(sorted(VALID_PROVIDER_TYPES))}",
        )
    if entity_type not in VALID_ENTITY_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="entity_type must be 'PF' or 'PJ'",
        )

    filename_lower = (file.filename or "").lower()
    content_type_ok = file.content_type == DOCX_MIME_TYPE
    extension_ok = filename_lower.endswith(".docx")
    if not content_type_ok and not extension_ok:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Only DOCX files are accepted for contract templates",
        )

    # Read content and check size (max 20 MB)
    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="File exceeds the 20 MB size limit",
        )

    # Verify the file is actually a DOCX/ZIP (magic bytes: PK\x03\x04).
    # This prevents someone from uploading a non-DOCX file with a .docx extension.
    if content[:4] != b"PK\x03\x04":
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="File content is not a valid DOCX document",
        )

    # Save to disk — fixed path: contract_templates/{provider_type}_{entity_type}.docx
    template_dir = await _get_template_dir(settings.DOCUMENTS_BASE_PATH)
    file_path = os.path.join(template_dir, f"{provider_type}_{entity_type}.docx")

    safe_original = sanitize_filename(
        file.filename or f"{provider_type}_{entity_type}_contrato.docx"
    )

    try:
        await asyncio.to_thread(_write_bytes_to_file, file_path, content)

        # Upsert in DB
        result = await db.execute(
            select(ContractTemplate).where(
                ContractTemplate.provider_type == provider_type,
                ContractTemplate.entity_type == entity_type,
            )
        )
        existing = result.scalar_one_or_none()

        now = datetime.now(timezone.utc)

        if existing:
            existing.file_path = file_path
            existing.original_filename = safe_original
            existing.uploaded_at = now
            existing.uploaded_by_analyst_id = current_analyst.id
            template_record = existing
        else:
            template_record = ContractTemplate(
                id=uuid.uuid4(),
                provider_type=provider_type,
                entity_type=entity_type,
                file_path=file_path,
                original_filename=safe_original,
                uploaded_at=now,
                uploaded_by_analyst_id=current_analyst.id,
            )
            db.add(template_record)

        await log_audit(
            db=db,
            action="contract_template_uploaded",
            analyst_id=current_analyst.id,
            metadata={
                "provider_type": provider_type,
                "entity_type": entity_type,
                "original_filename": safe_original,
                "analyst_email": current_analyst.email,
            },
        )
        await db.commit()
    except Exception:
        # Clean up the file if DB commit failed
        if await asyncio.to_thread(os.path.exists, file_path):
            await asyncio.to_thread(os.remove, file_path)
        raise

    await db.refresh(template_record, ["uploaded_by_analyst"])

    return ContractTemplateInfo(
        provider_type=template_record.provider_type,
        entity_type=template_record.entity_type,
        provider_type_label=PROVIDER_TYPE_LABELS.get(template_record.provider_type, template_record.provider_type),
        entity_type_label=ENTITY_TYPE_LABELS.get(template_record.entity_type, template_record.entity_type),
        original_filename=template_record.original_filename,
        uploaded_at=template_record.uploaded_at,
        uploaded_by_name=(
            template_record.uploaded_by_analyst.full_name
            if template_record.uploaded_by_analyst
            else None
        ),
    )



